"""문서 로더 모듈"""
import os
from pathlib import Path
from typing import List, Dict, Optional
import chardet
from loguru import logger
try:
    import pypdf
except ImportError:
    import PyPDF2 as pypdf

from docx import Document

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False
    logger.warning("openpyxl이 설치되지 않았습니다. Excel 파일 지원이 비활성화됩니다.")

import markdown
from src.Config.config import settings
from src.Utils.file_tracker import FileTracker


class DocumentLoader:
    """다양한 형식의 문서를 로드하는 클래스"""
    
    def __init__(self):
        self.supported_extensions = settings.supported_extensions
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # MB to bytes
        self.file_tracker = FileTracker()
    
    def load_documents(self, directory: Path = None, only_new: bool = True) -> List[Dict[str, str]]:
        """디렉토리의 모든 지원 문서 로드
        
        Args:
            directory: 로드할 디렉토리 (기본값: settings.input_dir)
            only_new: True면 신규/변경된 파일만, False면 모든 파일 (기본값: True)
        """
        if directory is None:
            directory = settings.input_dir
        
        # 삭제된 파일들의 메타데이터 정리
        self.file_tracker.remove_missing_files()
        
        documents = []
        all_documents = []
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    # 파일 크기 확인
                    if file_path.stat().st_size > self.max_file_size:
                        logger.warning(f"파일 크기 초과, 건너뜀: {file_path}")
                        continue
                    
                    content = self.load_file(file_path)
                    if content:
                        # input 디렉토리 기준 상대 경로 계산
                        try:
                            relative_path = file_path.relative_to(directory)
                        except ValueError:
                            # 상대 경로 계산 실패시 파일명만 사용
                            relative_path = file_path.name
                        
                        doc_info = {
                            'path': str(file_path),
                            'name': file_path.name,
                            'relative_path': str(relative_path),
                            'content': content,
                            'type': file_path.suffix.lower()
                        }
                        
                        all_documents.append(doc_info)
                        logger.debug(f"문서 스캔 완료: {file_path.name}")
                        
                except Exception as e:
                    logger.error(f"문서 로드 실패 {file_path}: {e}")
        
        # only_new가 True면 신규/변경된 파일만 필터링
        if only_new:
            documents = self.file_tracker.get_new_or_changed_files(all_documents)
            
            if len(documents) < len(all_documents):
                logger.info(f"전체 {len(all_documents)}개 파일 중 {len(documents)}개가 신규/변경됨")
                skipped_count = len(all_documents) - len(documents)
                logger.info(f"{skipped_count}개 파일은 이미 임베딩되어 건너뜀")
            else:
                logger.info(f"모든 {len(documents)}개 파일이 신규/변경됨")
        else:
            documents = all_documents
            logger.info(f"모든 파일 로드 모드: {len(documents)}개 파일")
        
        # 로드된 파일들을 임베딩 완료로 표시 (실제 임베딩은 RAG 서비스에서 처리)
        for doc in documents:
            logger.info(f"신규/변경 문서 로드: {doc['name']}")
        
        return documents
    
    def load_all_documents(self, directory: Path = None) -> List[Dict[str, str]]:
        """모든 문서를 강제로 로드 (임베딩 상태 무시)"""
        return self.load_documents(directory, only_new=False)
    
    def mark_documents_embedded(self, documents: List[Dict[str, str]]):
        """문서들을 임베딩 완료로 표시"""
        for doc in documents:
            file_path = Path(doc['path'])
            self.file_tracker.mark_file_embedded(file_path)
        logger.info(f"{len(documents)}개 문서를 임베딩 완료로 표시")
    
    def reset_embedding_status(self):
        """모든 임베딩 상태와 RAG 저장소 초기화"""
        import shutil
        
        # 1. 파일 추적 메타데이터 초기화
        self.file_tracker.clear_metadata()
        logger.info("모든 파일의 임베딩 상태 추적 초기화 완료")
        
        # 2. RAG 저장소 디렉토리 삭제
        storage_dir = settings.lightrag_working_dir
        if storage_dir.exists():
            try:
                shutil.rmtree(storage_dir)
                logger.info(f"RAG 저장소 디렉토리 삭제 완료: {storage_dir}")
            except Exception as e:
                logger.error(f"RAG 저장소 디렉토리 삭제 실패: {storage_dir}, 오류: {e}")
                # 필요한 경우 예외를 다시 발생시키거나 사용자에게 알림
                # raise e
        else:
            logger.info(f"RAG 저장소 디렉토리가 존재하지 않아 건너뜀: {storage_dir}")
    
    def get_embedding_status(self) -> Dict:
        """현재 임베딩 상태 정보 반환"""
        return self.file_tracker.get_status_info()

    def load_file(self, file_path: Path) -> Optional[str]:
        """파일 형식에 따라 적절한 로더 호출"""
        extension = file_path.suffix.lower()
        
        if extension == '.txt':
            return self._load_text(file_path)
        elif extension == '.pdf':
            return self._load_pdf(file_path)
        elif extension == '.docx':
            return self._load_docx(file_path)
        elif extension == '.md':
            return self._load_markdown(file_path)
        elif extension == '.xlsx':
            if EXCEL_SUPPORT:
                return self._load_excel(file_path)
            else:
                logger.warning(f"Excel 지원이 비활성화되어 파일을 건너뜁니다: {file_path}")
                return None
        else:
            logger.warning(f"지원하지 않는 파일 형식: {extension}")
            return None
    
    def _load_text(self, file_path: Path) -> str:
        """텍스트 파일 로드 (인코딩 자동 감지)"""
        # 인코딩 감지
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or settings.encoding
        
        # 파일 읽기
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # 감지 실패 시 기본 인코딩 시도
            with open(file_path, 'r', encoding=settings.encoding, errors='ignore') as f:
                return f.read()
    
    def _load_pdf(self, file_path: Path) -> str:
        """PDF 파일 로드"""
        text = []
        try:
            with open(file_path, 'rb') as f:
                # pypdf (최신) API 시도
                if hasattr(pypdf, 'PdfReader'):
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                # PyPDF2 (구버전) API 시도
                else:
                    reader = pypdf.PdfFileReader(f)
                    for page_num in range(reader.numPages):
                        page = reader.getPage(page_num)
                        page_text = page.extractText()
                        if page_text:
                            text.append(page_text)
        except Exception as e:
            logger.error(f"PDF 파일 로드 실패 {file_path}: {e}")
            return ""
        return '\n'.join(text)
    
    def _load_docx(self, file_path: Path) -> str:
        """DOCX 파일 로드"""
        doc = Document(str(file_path))
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    
    def _load_markdown(self, file_path: Path) -> str:
        """Markdown 파일 로드 및 텍스트 변환"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Markdown을 HTML로 변환 후 텍스트 추출
        html = markdown.markdown(md_content)
        # 간단한 HTML 태그 제거
        import re
        text = re.sub('<[^<]+?>', '', html)
        return text.strip()
    
    def _load_excel(self, file_path: Path) -> str:
        """Excel 파일 로드"""
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"## Sheet: {sheet_name}")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) if cell is not None else '' for cell in row]
                if any(row_text):  # 빈 행 제외
                    text.append(' | '.join(row_text))
        
        workbook.close()
        return '\n'.join(text)
    
    def get_document_stats(self, documents: List[Dict[str, str]]) -> Dict[str, any]:
        """문서 통계 반환"""
        stats = {
            'total_documents': len(documents),
            'total_characters': sum(len(doc['content']) for doc in documents),
            'by_type': {}
        }
        
        for doc in documents:
            doc_type = doc['type']
            if doc_type not in stats['by_type']:
                stats['by_type'][doc_type] = 0
            stats['by_type'][doc_type] += 1
        
        return stats 